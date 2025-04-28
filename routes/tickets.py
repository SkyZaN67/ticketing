from flask import Blueprint, render_template, request, redirect, url_for, send_file
from flask_login import login_required, current_user
from extensions import db, mail
from flask_mail import Message
from models import Ticket, Client, Utilisateur
from docx import Document
from io import BytesIO
from datetime import datetime

tickets_bp = Blueprint('tickets', __name__, url_prefix='/tickets')


def envoyer_mail_support(ticket):
    sujet = f"[SUPPORT] Nouveau ticket : {ticket.titre}"
    corps = f"""
Bonjour,

Un nouveau ticket a été assigné au support.

📌 Titre : {ticket.titre}
👤 Client : {ticket.client.nom}
📞 Téléphone : {ticket.telephone or 'Non renseigné'}
📝 Description : {ticket.description}
📣 Urgence : {ticket.urgence}
👨 Demandeur : {ticket.demandeur or 'Non précisé'}
🕒 Date de création : {ticket.date_creation.strftime('%d/%m/%Y %H:%M')}

Merci,
L’équipe IT360
    """
    msg = Message(subject=sujet, recipients=["support@is2s-st.com"])
    msg.body = corps
    mail.send(msg)


@tickets_bp.route('/create/<int:client_id>', methods=['GET', 'POST'])
@login_required
def create(client_id):
    utilisateurs = Utilisateur.query.filter_by(role='technicien').all()

    if request.method == 'POST':
        titre = request.form.get('titre')
        description = request.form.get('description')
        demandeur = request.form.get('demandeur')
        telephone = request.form.get('telephone')
        urgence = request.form.get('urgence')
        utilisateur_id = request.form.get('utilisateur_id')

        groupe_assignation = None
        if utilisateur_id == "support":
            groupe_assignation = "support"
            utilisateur_id = None
        else:
            utilisateur_id = int(utilisateur_id) if utilisateur_id else None

        ticket = Ticket(
            titre=titre,
            description=description,
            telephone=telephone,
            demandeur=demandeur,
            client_id=client_id,
            utilisateur_id=utilisateur_id,
            groupe_assignation=groupe_assignation,
            urgence=urgence
        )
        db.session.add(ticket)
        db.session.commit()

        if groupe_assignation == "support":
            envoyer_mail_support(ticket)

        return redirect(url_for('clients.detail', client_id=client_id))

    return render_template('create_ticket.html', client_id=client_id, utilisateurs=utilisateurs)


@tickets_bp.route('/create_by_technicien', methods=['GET', 'POST'])
@login_required
def create_by_technicien():
    if current_user.role != "technicien":
        return redirect(url_for('dashboard.index'))

    utilisateurs = Utilisateur.query.filter_by(role='technicien').all()
    clients = Client.query.all()

    if request.method == 'POST':
        titre = request.form.get('titre')
        description = request.form.get('description')
        demandeur = request.form.get('demandeur')
        telephone = request.form.get('telephone')
        urgence = request.form.get('urgence')
        client_id = request.form.get('client_id')
        utilisateur_id = request.form.get('utilisateur_id')

        groupe_assignation = None
        if utilisateur_id == "support":
            groupe_assignation = "support"
            utilisateur_id = None
        else:
            utilisateur_id = int(utilisateur_id) if utilisateur_id else None

        ticket = Ticket(
            titre=titre,
            description=description,
            telephone=telephone,
            demandeur=demandeur,
            client_id=int(client_id),
            utilisateur_id=utilisateur_id,
            groupe_assignation=groupe_assignation,
            urgence=urgence
        )
        db.session.add(ticket)
        db.session.commit()

        if groupe_assignation == "support":
            envoyer_mail_support(ticket)

        return redirect(url_for('dashboard.index'))

    return render_template('create_ticket_technicien.html', utilisateurs=utilisateurs, clients=clients)


@tickets_bp.route('/<int:ticket_id>', methods=['GET', 'POST'])
@login_required
def detail(ticket_id):
    ticket = Ticket.query.get_or_404(ticket_id)
    utilisateurs = Utilisateur.query.filter_by(role='technicien').all()

    if request.method == 'POST':
        if request.form.get('action') == 'self_assign' and current_user.role == "technicien":
            ticket.utilisateur_id = current_user.id
            ticket.groupe_assignation = None
            db.session.commit()
            return redirect(url_for('tickets.detail', ticket_id=ticket.id))

        ancien_statut = ticket.statut
        nouveau_statut = request.form.get('statut')
        ticket.statut = nouveau_statut

        if nouveau_statut == "En cours" and ticket.date_debut is None:
            ticket.date_debut = datetime.now()

        if ancien_statut != "Terminé" and nouveau_statut == "Terminé":
            if ticket.date_debut is None:
                ticket.date_debut = datetime.now()
            ticket.date_fin = datetime.now()

        # Temps passé
        mode_temps = request.form.get('mode_temps_passe')
        if mode_temps == 'manuel':
            temps_passe_form = request.form.get('temps_passe')
            ticket.temps_passe_personnalise = int(temps_passe_form) if temps_passe_form else None
        else:
            ticket.temps_passe_personnalise = None

        # Temps facturé
        temps_facture_perso = request.form.get('temps_facture_personnalise')
        ticket.temps_facture_personnalise = int(temps_facture_perso) if temps_facture_perso else None

        # Assignation
        assignation = request.form.get('utilisateur_id')
        if assignation == "support":
            ticket.utilisateur_id = None
            ticket.groupe_assignation = "support"
            envoyer_mail_support(ticket)
        else:
            ticket.utilisateur_id = int(assignation) if assignation else None
            ticket.groupe_assignation = None

        # Rapport, Observation, Matériel fourni
        ticket.rapport = request.form.get('rapport')
        ticket.observation = request.form.get('observation')
        ticket.materiel_fourni = 'materiel_fourni' in request.form

        db.session.commit()

        if current_user.role == 'technicien':
            return redirect(url_for('dashboard.index'))
        else:
            return redirect(url_for('clients.detail', client_id=ticket.client_id))

    return render_template('ticket_detail.html', ticket=ticket, utilisateurs=utilisateurs)



@tickets_bp.route('/delete/<int:ticket_id>', methods=['POST'])
@login_required
def delete(ticket_id):
    ticket = Ticket.query.get_or_404(ticket_id)
    client_id = ticket.client_id
    db.session.delete(ticket)
    db.session.commit()
    return redirect(url_for('clients.detail', client_id=client_id))


@tickets_bp.route('/export/<int:ticket_id>')
@login_required
def export_ticket(ticket_id):
    ticket = Ticket.query.get_or_404(ticket_id)
    doc = Document("Intervention Technique Vierge.docx")

    def format_minutes_hm(minutes):
        sign = '-' if minutes < 0 else ''
        minutes = abs(minutes)
        return f"{sign}{minutes // 60}h{minutes % 60:02d}"

    remplacements = {
        "{{CLIENT}}": ticket.client.nom,
        "{{EMAIL_CLIENT}}": ticket.client.email or "Non renseignée",
        "{{TECHNICIEN}}": ticket.utilisateur.nom if ticket.utilisateur else "Non assigné",
        "{{DATE}}": ticket.date_creation.strftime('%d/%m/%Y') if ticket.date_creation else "Non renseignée",
        "{{TEMPS}}": f"{ticket.temps_passe} min",
        "{{TEMPS_FACTURE}}": f"{format_minutes_hm(ticket.temps_facture)}",
        "{{DESCRIPTION}}": ticket.description or "Aucune description",
        "{{RAPPORT}}": ticket.rapport or "Aucun rapport",
        "{{OBSERVATION}}": ticket.observation or "Aucune observation",
        "{{MATERIEL_FOURNI}}": "Oui" if ticket.materiel_fourni else "Non",
        "{{STATUT}}": ticket.statut,
        "{{URGENCE}}": ticket.urgence or "Non définie",
        "{{DEMANDEUR}}": ticket.demandeur or "Non précisé"
    }

    def replace_placeholders(doc, replacements):
        for p in doc.paragraphs:
            for key, value in replacements.items():
                if key in p.text:
                    for run in p.runs:
                        if key in run.text:
                            run.text = run.text.replace(key, value)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        for key, value in replacements.items():
                            if key in p.text:
                                for run in p.runs:
                                    if key in run.text:
                                        run.text = run.text.replace(key, value)

    replace_placeholders(doc, remplacements)

    output = BytesIO()
    doc.save(output)
    output.seek(0)
    filename = f"fiche_intervention_ticket_{ticket.id}.docx"
    return send_file(output, as_attachment=True, download_name=filename)

